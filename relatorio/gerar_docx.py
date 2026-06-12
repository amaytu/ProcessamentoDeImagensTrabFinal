"""
gerar_docx.py — Gera o relatório científico em formato .docx
Uso: python relatorio/gerar_docx.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "relatorio.docx")


def set_cell_shading(cell, color_hex: str):
    """Aplica cor de fundo a uma célula de tabela."""
    shading = cell._element.get_or_add_tcPr()
    sh = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(sh)


def add_styled_paragraph(doc, text, style="Normal", bold=False, italic=False,
                         alignment=None, font_size=None, space_after=None,
                         space_before=None, font_name=None, color=None,
                         first_line_indent=None):
    """Adiciona um parágrafo com formatação customizada."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p


def add_mixed_paragraph(doc, parts, alignment=None, space_after=None,
                        space_before=None, first_line_indent=None,
                        line_spacing=None):
    """Adiciona parágrafo com múltiplos runs (bold/italic misturados)."""
    p = doc.add_paragraph()
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    if line_spacing is not None:
        p.paragraph_format.line_spacing = Pt(line_spacing)
    return p


def body_text(doc, text, first_indent=True):
    """Parágrafo de corpo de texto padrão (justificado, Times 12, espaçamento 1.5)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)  # 1.5 lines
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    return p


def heading(doc, text, level=1):
    """Adiciona heading numerado."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(6)
    return h


def add_table(doc, headers, rows):
    """Adiciona uma tabela formatada."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"

    doc.add_paragraph()  # spacing after table
    return table


def build_document():
    doc = Document()

    # ─── Page margins ───
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # ─── Default font ───
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # ═══════════════════════════════════════════════════════════════════
    # TÍTULO
    # ═══════════════════════════════════════════════════════════════════
    add_styled_paragraph(
        doc,
        "Análise Biomecânica de Rosca Direta em Tempo Real\n"
        "Utilizando Visão Computacional e Estimativa de Pose",
        bold=True, font_size=16, font_name="Times New Roman",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12, space_before=24,
    )

    # Autores
    add_styled_paragraph(
        doc,
        "Gabriel [Sobrenome]¹, [Membro 2]¹, [Membro 3]¹",
        bold=True, font_size=12, font_name="Times New Roman",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )

    add_styled_paragraph(
        doc,
        "¹ Universidade Regional de Blumenau (FURB) – Blumenau, SC – Brasil",
        font_size=10, font_name="Times New Roman",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )

    add_styled_paragraph(
        doc,
        "Disciplina: Processamento de Imagens — Prof. Aurélio Hoppe",
        font_size=10, font_name="Times New Roman",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )

    add_styled_paragraph(
        doc,
        "Junho de 2026",
        font_size=10, font_name="Times New Roman",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=18,
    )

    # ═══════════════════════════════════════════════════════════════════
    # RESUMO
    # ═══════════════════════════════════════════════════════════════════
    add_styled_paragraph(
        doc, "Resumo", bold=True, font_size=12,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6, space_before=12,
    )

    body_text(
        doc,
        "Este trabalho apresenta o desenvolvimento de um sistema de visão computacional "
        "em tempo real para análise biomecânica do exercício de rosca direta (biceps curl). "
        "O sistema utiliza o framework MediaPipe Pose, baseado na arquitetura BlazePose, para "
        "estimar a pose corporal do praticante a partir de um fluxo de vídeo (webcam ou arquivo). "
        "A partir dos landmarks articulares detectados — ombro, cotovelo e pulso — o ângulo "
        "interno do cotovelo é calculado quadro a quadro por meio da função arco-tangente de "
        "dois argumentos (arctan2), garantindo estabilidade numérica em toda a faixa angular "
        "de 0° a 180°. Uma máquina de estados finitos classifica automaticamente as fases do "
        "movimento (excêntrica e concêntrica), contabiliza repetições válidas e detecta "
        "execuções com amplitude insuficiente, caracterizadas como \"roubo\" biomecânico. "
        "O sistema alcança taxas de processamento superiores a 30 FPS em hardware de consumo "
        "geral sem necessidade de GPU dedicada, demonstrando a viabilidade de ferramentas "
        "acessíveis para monitoramento de qualidade de execução em treinamento resistido.",
        first_indent=False,
    )

    add_mixed_paragraph(
        doc,
        [
            ("Palavras-chave: ", True, False),
            ("Visão computacional. Estimativa de pose. Biomecânica. "
             "Rosca direta. MediaPipe. Processamento de imagens.", False, False),
        ],
        space_after=12,
    )

    # ═══════════════════════════════════════════════════════════════════
    # 1. INTRODUÇÃO
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "1. Introdução", level=1)

    body_text(
        doc,
        "O treinamento resistido é amplamente reconhecido como um dos pilares para o "
        "desenvolvimento de força e hipertrofia muscular. Dentre os mecanismos que governam "
        "a adaptação muscular, a tensão mecânica destaca-se como o principal estímulo "
        "fisiológico para a hipertrofia (SCHOENFELD, 2010). A qualidade da execução dos "
        "exercícios — particularmente a amplitude de movimento (Range of Motion, ROM) — é um "
        "fator determinante para maximizar a tensão mecânica imposta ao músculo-alvo. Estudos "
        "recentes demonstram que o treinamento com amplitude total, enfatizando a posição "
        "alongada do músculo, promove adaptações hipertróficas superiores em comparação ao "
        "treino com amplitude parcial na porção encurtada (PEDROSA et al., 2023).",
    )

    body_text(
        doc,
        "Apesar da importância da execução correta, a avaliação biomecânica de exercícios "
        "tradicionalmente requer equipamentos laboratoriais especializados (sistemas de captura "
        "de movimento ópticos, plataformas de força) ou a supervisão presencial de um "
        "profissional de educação física. Neste contexto, avanços recentes em visão "
        "computacional e estimativa de pose humana (Human Pose Estimation, HPE) oferecem "
        "alternativas acessíveis e não invasivas para monitoramento em tempo real. Frameworks "
        "como o MediaPipe, desenvolvido pelo Google, disponibilizam modelos de redes neurais "
        "leves capazes de inferir a posição de 33 pontos articulares (landmarks) do corpo "
        "humano a partir de uma única câmera RGB, operando em tempo real em dispositivos "
        "convencionais (BAZAREVSKY et al., 2020; LUGARESI et al., 2019).",
    )

    body_text(
        doc,
        "O presente trabalho propõe o desenvolvimento de uma ferramenta de software que "
        "integra técnicas de processamento de imagens, estimativa de pose e análise "
        "biomecânica para monitorar a execução do exercício de rosca direta (biceps curl). "
        "O sistema classifica as fases do movimento, conta repetições válidas e emite alertas "
        "visuais quando a amplitude do movimento é insuficiente, auxiliando o praticante na "
        "maximização do estímulo de tensão mecânica.",
    )

    # ═══════════════════════════════════════════════════════════════════
    # 2. FUNDAMENTAÇÃO TEÓRICA
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "2. Fundamentação Teórica", level=1)

    # 2.1
    heading(doc, "2.1. Estimativa de Pose Humana com BlazePose", level=2)

    body_text(
        doc,
        "A estimativa de pose humana é uma tarefa fundamental em visão computacional que "
        "consiste em localizar as articulações do corpo humano a partir de imagens ou vídeos. "
        "A arquitetura BlazePose, proposta por Bazarevsky et al. (2020), foi projetada "
        "especificamente para operação em tempo real em dispositivos com recursos "
        "computacionais limitados. O modelo emprega um pipeline de duas etapas: (i) um "
        "detector de pose baseado em rede neural convolucional (CNN) leve que localiza a "
        "região de interesse (ROI) do corpo humano, e (ii) um rastreador de landmarks que "
        "prediz as coordenadas 2D/3D de 33 pontos articulares dentro da ROI detectada.",
    )

    body_text(
        doc,
        "O BlazePose utiliza uma abordagem híbrida que combina mapas de calor (heatmaps) "
        "com regressão direta de coordenadas, alcançando um equilíbrio entre precisão e "
        "velocidade de inferência. O modelo é capaz de processar mais de 30 quadros por "
        "segundo em dispositivos móveis com processadores convencionais, eliminando a "
        "dependência de unidades de processamento gráfico (GPUs) dedicadas para a "
        "inferência local.",
    )

    body_text(
        doc,
        "O framework MediaPipe (LUGARESI et al., 2019) encapsula o BlazePose em uma "
        "arquitetura de grafos modular que permite a construção de pipelines completos de "
        "processamento de mídia — desde a captura de vídeo até a inferência e renderização "
        "de resultados — de forma eficiente e multiplataforma.",
    )

    # 2.2
    heading(doc, "2.2. Cálculo Angular e Estabilidade Numérica", level=2)

    body_text(
        doc,
        "Para quantificar o movimento articular, o sistema calcula o ângulo interno do "
        "cotovelo formado pelos vetores ombro→cotovelo (BA) e pulso→cotovelo (BC), onde B "
        "é a posição do cotovelo. O método escolhido utiliza a função arco-tangente de dois "
        "argumentos (arctan2):",
    )

    add_styled_paragraph(
        doc,
        "θ = |arctan2(BA × BC, BA · BC)|",
        font_size=12, font_name="Times New Roman", italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6, space_before=6,
    )

    body_text(
        doc,
        "onde × denota o produto vetorial 2D (componente escalar z) e · denota o produto "
        "escalar. Esta formulação é numericamente superior à alternativa baseada em arccos, "
        "que apresenta derivadas tendendo ao infinito nos extremos (0° e 180°), resultando "
        "em instabilidade numérica e sensibilidade a ruídos de medição nessas regiões "
        "críticas para a análise do exercício (BRADSKI, 2000).",
    )

    # 2.3
    heading(doc, "2.3. Tensão Mecânica e Amplitude de Movimento", level=2)

    body_text(
        doc,
        "A tensão mecânica é definida como a força gerada nas fibras musculares durante a "
        "contração sob carga, sendo considerada o principal estímulo para a sinalização de "
        "vias anabólicas que culminam na hipertrofia muscular (SCHOENFELD, 2010). Schoenfeld "
        "identificou que a combinação de alta magnitude de força com duração suficiente de "
        "tensão — particularmente durante o alongamento ativo do músculo sob carga — "
        "maximiza a resposta adaptativa.",
    )

    body_text(
        doc,
        "Pedrosa et al. (2023) investigaram especificamente o efeito da amplitude de "
        "movimento na hipertrofia do bíceps braquial durante o exercício de rosca preacher. "
        "Os resultados demonstraram que o treino na amplitude inicial (posição alongada, "
        "0°–68°) promoveu ganhos de área de secção transversa significativamente superiores "
        "em comparação ao treino na amplitude final (posição encurtada, 68°–135°). Estes "
        "achados fundamentam a decisão de design do sistema em alertar o usuário quando a "
        "fase excêntrica (descida do peso com o braço estendendo) é interrompida prematuramente "
        "— isto é, quando a extensão máxima atingida é inferior a 150°, indicando que o "
        "praticante não está explorando adequadamente a amplitude na posição alongada.",
    )

    # ═══════════════════════════════════════════════════════════════════
    # 3. METODOLOGIA
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "3. Metodologia", level=1)

    # 3.1
    heading(doc, "3.1. Arquitetura do Sistema", level=2)

    body_text(
        doc,
        "O sistema foi implementado em Python 3.10+ seguindo princípios de modularidade e "
        "separação de responsabilidades. A arquitetura é composta por quatro módulos "
        "principais, conforme apresentado na Tabela 1.",
    )

    add_table(doc,
        ["Módulo", "Responsabilidade"],
        [
            ["main.py", "Loop principal, captura de vídeo (OpenCV), integração com MediaPipe PoseLandmarker e orquestração do pipeline"],
            ["biomechanics.py", "Cálculo do ângulo articular via arctan2 e extração de coordenadas de landmarks"],
            ["exercise_analyzer.py", "Máquina de estados finitos para classificação de fases, contagem de repetições e detecção de amplitude insuficiente"],
            ["hud_renderer.py", "Renderização do HUD (Head-Up Display) sobreposto ao vídeo: esqueleto, ângulo em tempo real, barra de progresso e alertas visuais"],
        ],
    )

    add_styled_paragraph(
        doc, "Tabela 1 — Módulos do sistema e suas responsabilidades.",
        font_size=10, italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
    )

    # 3.2
    heading(doc, "3.2. Pipeline de Processamento", level=2)

    body_text(
        doc,
        "O fluxo de processamento para cada quadro de vídeo é executado sequencialmente:",
        first_indent=False,
    )

    steps = [
        "Captura: leitura do quadro via OpenCV (VideoCapture), com suporte a webcam e arquivos de vídeo.",
        "Pré-processamento: conversão de espaço de cores BGR→RGB e espelhamento horizontal (modo webcam).",
        "Inferência de Pose: detecção de 33 landmarks corporais via PoseLandmarker no modo VIDEO, com limiar de confiança de detecção de 70% e de rastreamento de 50%.",
        "Cálculo Biomecânico: extração das coordenadas normalizadas de ombro, cotovelo e pulso, seguida do cálculo angular por arctan2.",
        "Análise de Exercício: atualização da máquina de estados com o ângulo atual, retornando fase, contagem de repetições e status de alerta.",
        "Renderização do HUD: sobreposição gráfica com esqueleto articular, ângulo numérico, painel de informações, barra de progresso e banner de alerta pulsante.",
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(step)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

    doc.add_paragraph()  # spacing

    # 3.3
    heading(doc, "3.3. Máquina de Estados para Contagem de Repetições", level=2)

    body_text(
        doc,
        "A lógica de análise do exercício é modelada como uma máquina de estados finitos "
        "determinística com três estados:",
        first_indent=False,
    )

    states = [
        ("IDLE: ", "estado inicial — aguarda extensão completa do braço (ângulo ≥ 160°) para iniciar a análise."),
        ("ECCENTRIC (fase excêntrica): ", "braço em descida — monitora o ângulo máximo atingido e transiciona para CONCENTRIC quando o ângulo atinge ≤ 40°. Se o ângulo máximo registrado na fase for inferior a 150°, um alerta de amplitude insuficiente é disparado."),
        ("CONCENTRIC (fase concêntrica): ", "braço em subida — ao atingir novamente ≥ 160°, uma repetição válida é contabilizada e o estado retorna para ECCENTRIC."),
    ]
    for label, desc in states:
        p = doc.add_paragraph(style="List Bullet")
        run_b = p.add_run(label)
        run_b.bold = True
        run_b.font.size = Pt(12)
        run_b.font.name = "Times New Roman"
        run_d = p.add_run(desc)
        run_d.font.size = Pt(12)
        run_d.font.name = "Times New Roman"

    body_text(
        doc,
        "Uma repetição só é contada quando o ciclo completo extensão → contração → extensão "
        "é verificado, garantindo que apenas movimentos com amplitude biomecânica adequada "
        "sejam validados.",
    )

    # ═══════════════════════════════════════════════════════════════════
    # 4. RESULTADOS E DISCUSSÃO
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "4. Resultados e Discussão", level=1)

    # 4.1
    heading(doc, "4.1. Desempenho em Tempo Real", level=2)

    body_text(
        doc,
        "O sistema foi avaliado em um computador pessoal com processador convencional "
        "(sem GPU dedicada), atingindo consistentemente taxas de processamento superiores "
        "a 30 FPS com resolução de captura de 1280×720 pixels. Este resultado confirma a "
        "viabilidade do modelo BlazePose Lite para aplicações em tempo real, conforme "
        "reportado por Bazarevsky et al. (2020).",
    )

    # 4.2
    heading(doc, "4.2. Precisão do Cálculo Angular", level=2)

    body_text(
        doc,
        "A acurácia do cálculo angular foi validada por meio de testes automatizados com "
        "vetores de referência em ângulos conhecidos (0°, 45°, 60°, 90°, 120° e 180°), "
        "todos com erro absoluto inferior a 0,1°. A escolha da função arctan2 se mostrou "
        "essencial para manter a estabilidade nos extremos da faixa angular, onde a "
        "formulação baseada em arccos apresentaria instabilidade numérica.",
    )

    # 4.3
    heading(doc, "4.3. Validação da Máquina de Estados", level=2)

    body_text(
        doc,
        "A máquina de estados foi testada com sequências simuladas de ângulos que reproduzem "
        "cenários de execução correta, repetições consecutivas e execução com amplitude "
        "insuficiente (\"roubo\"). Todos os cenários de teste foram validados com sucesso, "
        "incluindo:",
        first_indent=False,
    )

    validations = [
        "Transição correta entre estados (IDLE → ECCENTRIC → CONCENTRIC → ECCENTRIC).",
        "Contagem incremental apenas ao completar o ciclo biomecânico.",
        "Disparo e expiração do alerta de amplitude insuficiente.",
    ]
    for v in validations:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(v)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

    # 4.4
    heading(doc, "4.4. Limitações", level=2)

    body_text(
        doc,
        "O sistema apresenta algumas limitações: (i) a estimativa de pose é realizada em "
        "2D, desconsiderando a profundidade do movimento; (ii) a análise é restrita ao plano "
        "frontal da câmera; (iii) o modelo é sensível a oclusões parciais dos landmarks "
        "articulares; e (iv) os limiares de ângulo são fixos e podem não ser ideais para "
        "todas as anatomias.",
    )

    # ═══════════════════════════════════════════════════════════════════
    # 5. CONCLUSÃO
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "5. Conclusão", level=1)

    body_text(
        doc,
        "Este trabalho demonstrou a viabilidade de um sistema acessível de análise "
        "biomecânica em tempo real para o exercício de rosca direta, utilizando "
        "exclusivamente uma câmera RGB e técnicas de processamento de imagens baseadas em "
        "aprendizado profundo. A integração do MediaPipe Pose com cálculos de cinemática "
        "articular via arctan2 e uma máquina de estados finitos possibilitou a classificação "
        "automática de fases, a contagem precisa de repetições e a detecção de amplitude "
        "insuficiente — contribuindo para a conscientização do praticante sobre a importância "
        "da amplitude total na maximização da tensão mecânica.",
    )

    body_text(
        doc,
        "Como trabalhos futuros, propõe-se a extensão da análise para o espaço 3D "
        "utilizando as coordenadas tridimensionais fornecidas pelo BlazePose, a "
        "implementação de limiares adaptativos baseados na anatomia do usuário, o suporte "
        "a múltiplos exercícios e a integração com plataformas de treinamento para registro "
        "longitudinal de desempenho.",
    )

    # ═══════════════════════════════════════════════════════════════════
    # REFERÊNCIAS
    # ═══════════════════════════════════════════════════════════════════
    heading(doc, "Referências", level=1)

    references = [
        "BAZAREVSKY, V.; GRISHCHENKO, I.; RAVEENDRAN, K.; ZHU, T.; ZHANG, F.; "
        "GRUNDMANN, M. BlazePose: On-device Real-time Body Pose tracking. In: CVPR "
        "Workshop on Computer Vision for Augmented and Virtual Reality, 2020. Disponível "
        "em: https://arxiv.org/abs/2006.10204. Acesso em: 12 jun. 2026.",

        "BRADSKI, G. The OpenCV Library. Dr. Dobb's Journal of Software Tools, v. 25, "
        "n. 11, p. 120–125, 2000.",

        "HARRIS, C. R.; MILLMAN, K. J.; VAN DER WALT, S. J. et al. Array programming "
        "with NumPy. Nature, v. 585, n. 7825, p. 357–362, 2020. "
        "DOI: 10.1038/s41586-020-2649-2.",

        "LUGARESI, C.; TANG, J.; NASH, H.; MCCLANAHAN, C.; UBOWEJA, E.; HAYS, M.; "
        "ZHANG, F.; CHANG, C.-L.; YONG, M. G.; LEE, J.; CHANG, W.-T.; HUA, W.; GEORG, M.; "
        "GRUNDMANN, M. MediaPipe: A Framework for Building Perception Pipelines. arXiv "
        "preprint, 2019. Disponível em: https://arxiv.org/abs/1906.08172. Acesso em: "
        "12 jun. 2026.",

        "PEDROSA, G. F.; LIMA, F. V.; SCHOENFELD, B. J.; LACERDA, L. T.; SIMÕES, M. G.; "
        "PEREIRA, M. R.; DINIZ, R. C. R.; CHAGAS, M. H. Training in the Initial Range of "
        "Motion Promotes Greater Muscle Adaptations Than at Final in the Arm Curl. Sports, "
        "v. 11, n. 3, p. 55, 2023. DOI: 10.3390/sports11030055.",

        "SCHOENFELD, B. J. The Mechanisms of Muscle Hypertrophy and Their Application to "
        "Resistance Training. Journal of Strength and Conditioning Research, v. 24, n. 10, "
        "p. 2857–2872, 2010. DOI: 10.1519/JSC.0b013e3181e840f3.",
    ]

    for ref in references:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-1.25)  # hanging indent

    # ─── Save ───
    doc.save(OUTPUT)
    print(f"[OK] Relatório salvo em: {OUTPUT}")


if __name__ == "__main__":
    build_document()
